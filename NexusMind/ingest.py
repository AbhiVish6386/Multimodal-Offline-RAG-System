# ingest.py
import os
from PIL import Image
import pytesseract
import docx
import whisper
import fitz  # pymupdf
import tempfile
import torch
import re

# Use Whisper-medium to balance accuracy and performance
DEFAULT_WHISPER_MODEL = "medium"

def _normalize_text(text):
    """Normalize extracted text across sources for cleaner embeddings.

    - Remove non-printable/hidden characters
    - Collapse excessive whitespace/newlines
    - Trim each line and drop empty lines
    """
    if not text:
        return ""
    # Remove problematic unicode control characters
    text = re.sub(r"[\u0000-\u0008\u000B\u000C\u000E-\u001F\u007F]", " ", text)
    # Normalize line endings
    text = text.replace("\r\n", "\n").replace("\r", "\n")
    # Trim lines and drop empties
    lines = [ln.strip() for ln in text.split("\n")]
    lines = [ln for ln in lines if ln]
    # Collapse multiple internal spaces
    lines = [re.sub(r"\s+", " ", ln) for ln in lines]
    # Re-join with single newlines to preserve some structure
    cleaned = "\n".join(lines)
    return cleaned.strip()

def extract_text_pdf(path):
    # Robust PDF extraction using PyMuPDF; return per-page chunks
    pages = []
    try:
        doc = fitz.open(path)
        for page_index, page in enumerate(doc, start=1):
            txt = page.get_text().strip()
            txt = _normalize_text(txt)
            if txt:
                pages.append((page_index, txt))
        return pages
    except Exception as e:
        print("PDF extract error:", e)
        return []

def extract_text_docx(path):
    try:
        d = docx.Document(path)

        parts = []
        # Collect paragraphs, including list-like paragraphs
        for p in d.paragraphs:
            txt = p.text.strip()
            if not txt:
                continue
            # Heuristic: add a bullet prefix if paragraph seems to be a list item
            # python-docx doesn't expose list type directly; detect leading bullets/dashes
            # bullet_like = bool(re.match(r"^[\u2022\-•\*]\\s+", txt))
            # if bullet_like:
            #     parts.append(txt)
            # else:
            #     parts.append(txt)
            parts.append(_normalize_text(txt))

        # Collect table cell contents row-wise
        for table in d.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    cell_txt = _normalize_text(cell.text)
                    if cell_txt:
                        row_text.append(cell_txt)
                if row_text:
                    parts.append(" \u2502 ".join(row_text))

        full_text = "\n".join(parts)
        return _normalize_text(full_text)
    except Exception as e:
        print("DOCX error:", e)
        return ""

def extract_text_image(path):
    try:
        img = Image.open(path)
        text = pytesseract.image_to_string(img)
        return _normalize_text(text)
    except Exception as e:
        print("OCR error:", e)
        return ""

def transcribe_audio(path, model_name=None):
    model_name = model_name or DEFAULT_WHISPER_MODEL
    try:
        device = "cuda" if torch.cuda.is_available() else "cpu"
        model = whisper.load_model(model_name, device=device)  # balanced accuracy/perf
        res = model.transcribe(path)
        # returns full text and word timestamps in "segments"
        return res.get("text", ""), res.get("segments", [])
    except Exception as e:
        print("Whisper error:", e)
        return "", []

def extract_text_txt(path):
    try:
        with open(path, 'r', encoding='utf-8') as f:
            return _normalize_text(f.read())
    except Exception as e:
        print("TXT error:", e)
        return ""

# helper to auto-detect filetype
def ingest_file(path):
    ext = os.path.splitext(path)[1].lower()
    chunks = []
    
    if ext == ".pdf":
        pages = extract_text_pdf(path)
        for page_num, page_text in pages:
            if page_text.strip():
                chunks.append((page_text, {"type":"pdf", "path":path, "page": page_num}))
    elif ext == ".docx":
        txt = extract_text_docx(path)
        if txt.strip():
            chunks.append((txt, {"type":"docx", "path":path}))
    elif ext == ".txt":
        txt = extract_text_txt(path)
        if txt.strip():
            chunks.append((txt, {"type":"txt", "path":path}))
    elif ext in [".png",".jpg",".jpeg",".bmp","tiff"]:
        txt = extract_text_image(path)
        if txt.strip():
            chunks.append((txt, {"type":"image", "path":path}))
    elif ext in [".mp3",".wav",".m4a"]:
        txt, segments = transcribe_audio(path)
        if txt.strip():
            # For audio files, we can split into chunks based on segments
            if segments:
                for i, segment in enumerate(segments):
                    start_time = segment.get('start', 0)
                    end_time = segment.get('end', 0)
                    text = segment.get('text', '').strip()
                    if text:
                        metadata = {
                            "type": "audio",
                            "path": path,
                            "start_time": f"{int(start_time//60):02d}:{int(start_time%60):02d}",
                            "end_time": f"{int(end_time//60):02d}:{int(end_time%60):02d}",
                            "segment_id": i
                        }
                        chunks.append((text, metadata))
            else:
                # Fallback to single chunk if no segments
                chunks.append((txt, {"type":"audio", "path":path}))
    
    return chunks, {"type":"unknown","path":path} if not chunks else chunks[0][1]
